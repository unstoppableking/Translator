# -*- coding: utf-8 -*-
"""
Created on Sat Oct  5 09:50:38 2019
读取docx文件，并用谷歌翻译将文件按段落翻译完并将翻译内容保存在新的文件里面
@author: XL
"""
import os
from multiprocessing import Pool
from docx import Document
from Translator import Translator

def read_file(file_name):
    document = Document(file_name)
    file_content = []
    for con in document.paragraphs:
        file_content.append(con.text)
    return file_content

def save_file(content):
    document = Document()
    for p_n in range(len(content)):
        document.add_paragraph(content[p_n])
    document.save(content[0][0:5]+'.docx') #保存文档

def translate(dict_content):
    content = {}
    translator = Translator()
    for index,con in dict_content.items():
        if len(con) < 7:
            content[index] = con
        else:
            con1 = translator.translate(con)
#            print('size = ', len(con1))
            content[index] = con1
    return content


def main():
    path = r'C:\Users\XL\Documents\GitHub\GoogleTranslator'.replace('\\', '/')
    file_list = ['t_3. Manuscript.docx', 't_4. Manuscript.docx']
    for f in file_list:
        file = os.path.join(path, f)
        file_content = read_file(file)
        index = range(0, len(file_content))
        dict_content = dict(zip(index, file_content))
#        pool = Pool(4)
#        pool.map(translate, dict_content)
#        pool.close()
#        pool.join()
        content = translate(dict_content)
        save_file(content)

if __name__ == '__main__':
    main()