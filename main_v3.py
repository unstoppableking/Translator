# -*- coding: utf-8 -*-
"""
Created on Sat Oct  5 09:50:38 2019
更新概况:
1.读取docx文件，将翻译完的内容和英文内容一起保存在新文件中;
2.开启多线程,但效果不是很明显;
3.不兼容前面版本
4.增加显示进度
@author: XL
"""
import os
from multiprocessing import Pool
from docx import Document
from docx.oxml.ns import qn
from Translator import Translator

def read_file(file_name):
    document = Document(file_name)
    file_content = []
    for con in document.paragraphs:
        file_content.append(con.text)
    return file_content


# 此处为v2版本改动函数,但兼容上一版本的此函数
def save_file(content, dict_content=None):
    document = Document()
    document.styles['Normal'].font.name = u'宋体'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    document.styles['Body Text'].font.name = u'Time New Roman'
    document.styles['Body Text']._element.rPr.rFonts.set(qn('w:english'), u'Time New Roman')
    hot = True
    if dict_content==None:
        hot = False
    for p_n in range(len(content)):
        if hot:
            document.add_paragraph(dict_content[p_n], style='Body Text')
        document.add_paragraph(content[p_n])
    document.save(content[0][0:5]+'.docx')  # 保存文档


# 此处为v3版本改动函数,不兼容上一版本的此函数
def translate(dict_content, translator, num):
    content = {}
    for index, con in dict_content.items():
        if len(con) < 7:
            content[index] = con
        else:
            con1 = translator.translate(con)
            print('Translate para %s/%s' % (index, num))
            content[index] = con1
    return content


# 此处为v3版本改动,不兼容上一版本函数
def main(mulp=False):
    path = r'C:\Users\XL\Documents\GitHub\GoogleTranslator'.replace('\\', '/')
    file_list = ['t_3. Manuscript.docx', 't_4. Manuscript.docx']
    translator = Translator()
    for f in file_list:
        file = os.path.join(path, f)
        file_content = read_file(file)
        index = range(0, len(file_content))
        dict_content = dict(zip(index, file_content))
        # fun_temp = lambda dict_c: translate(dict_c, translator)
        if mulp:
            content = {}
            pool = Pool(10)
            for i, j in dict_content.items():
                # a = pool.apply_async(translate, ({i: j}, translator)).get()
                content.update(pool.apply_async(translate, ({i: j}, translator, len(dict_content))).get())
            # temp = pool.map(fun_temp, dict_content)
            # for di in temp:
            #     content.update(di)
            # content.update(pool.map(fun_temp, dict_content))
            pool.close()
            pool.join()
        else:
            content = translate(dict_content, translator)
        save_file(content, dict_content)


if __name__ == '__main__':
    main(mulp=True)
