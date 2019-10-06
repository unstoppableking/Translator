# -*- coding: utf-8 -*-
"""
Created on Sat Oct  5 09:50:38 2019
1.读取docx文件，将翻译完的内容和英文内容一起保存在新文件中
2.增加显示进度
@author: XL
"""
import os
from docx import Document
from Translator import Translator


def read_file(file_name):
    document = Document(file_name)
    file_content = []
    for con in document.paragraphs:
        file_content.append(con.text)
    return file_content


# 相比前版本改动此处函数,但兼容上一版本的此函数
def save_file(content, dict_content=None):
    document = Document()
    hot = True
    if dict_content==None:
        hot = False
    for p_n in range(len(content)):
        if hot:
            document.add_paragraph(dict_content[p_n])
        document.add_paragraph(content[p_n])
    document.save(content[0][0:5]+'.docx') #保存文档


def translate(dict_content):
    content = {}
    translator = Translator()
    for index, con in dict_content.items():
        if len(con) < 7:
            content[index] = con
        else:
            con1 = translator.translate(con)
            print('Translate para %s/%s' %(index, len(dict_content)))
            content[index] = con1
    return content


def main():
    path = r'C:\Users\XL\Documents\GitHub\GoogleTranslator'.replace('\\', '/')
    file_list = ['t_3. Manuscript.docx', 't_4. Manuscript.docx']
    for f in file_list:
        file = os.path.join(path, f)
        file_content = read_file(file)
        index = range(0, len(file_content))
        dict_content = dict(zip(index, file_content))
#        pool = Pool(4)
#        pool.map(translate, dict_content)
#        pool.close()
#        pool.join()
        content = translate(dict_content)
        save_file(content, dict_content)

if __name__ == '__main__':
    main()
